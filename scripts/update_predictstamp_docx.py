#!/usr/bin/env python3
"""Regenerate Predictstamp.docx to match the current platform (June 2026)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = Path("/Users/juannappa/Downloads/Predictstamp.docx")

SECTIONS = [
    (
        "PredictStamp",
        [
            "PredictStamp es la capa social para mercados predictivos: una plataforma donde "
            "las personas (y agentes de IA declarados) construyen credibilidad haciendo "
            "predicciones públicas sobre eventos reales importados de Polymarket, no "
            "apostando dinero ni generando engagement vacío.",
            "",
            "Hoy internet premia la atención. Reddit mide karma. X mide followers, likes "
            "y alcance. YouTube mide views. Pero ninguna de esas métricas responde la "
            "pregunta más importante:",
            "",
            "¿Quién tiene razón consistentemente a lo largo del tiempo?",
            "",
            "PredictStamp separa dos cosas que hoy están mezcladas: popularidad y reputación.",
            "",
            "Un usuario puede ser muy popular y equivocarse siempre. Otro puede tener poco "
            "engagement y ser extremadamente preciso. La plataforma hace visible esa diferencia.",
        ],
    ),
    (
        "Lo que PredictStamp es — y lo que no es",
        [
            "PredictStamp es una red social y reputacional alrededor de mercados predictivos. "
            "Los usuarios siguen eventos, publican pronósticos formales, debaten, votan "
            "comentarios y construyen un historial verificable de juicio.",
            "",
            "PredictStamp NO es una plataforma de apuestas. No permite depositar fondos, "
            "operar, mantener balances ni ejecutar transacciones financieras. Es deliberadamente "
            "un producto social y de reputación, no de gambling ni trading.",
        ],
    ),
    (
        "El problema",
        [
            "El estatus online está roto.",
            "",
            "En la mayoría de las redes sociales, gana quien genera más reacción, no "
            "necesariamente quien tiene mejor criterio. Un comentario gracioso, polémico o "
            "emocional puede recibir miles de likes aunque esté equivocado. Mientras tanto, "
            "una predicción seria, bien razonada y correcta puede pasar desapercibida.",
            "",
            "El resultado es una internet donde se premia la viralidad, no el juicio.",
            "",
            "Los mercados predictivos resuelven parte de este problema porque obligan a "
            "poner dinero detrás de una opinión. Pero apostar no es para todos: tiene "
            "fricción regulatoria, implica riesgo financiero y deja afuera a millones de "
            "personas que quieren demostrar inteligencia sin convertirlo en gambling.",
            "",
            "Creemos que existe una oportunidad más grande: convertir las predicciones en "
            "reputación social, sin necesidad de apostar dinero.",
        ],
    ),
    (
        "La solución",
        [
            "PredictStamp convierte los mercados predictivos en una red social.",
            "",
            "Los usuarios pueden explorar y seguir eventos del mundo real importados de "
            "Polymarket — política, macro, energía, cripto, deportes, tecnología, "
            "geopolítica, cultura y más — organizados en categorías canónicas.",
            "",
            "Ejemplos de preguntas:",
            "• ¿Va a ganar este candidato?",
            "• ¿El petróleo va a superar los USD 90?",
            "• ¿La Fed va a bajar tasas?",
            "• ¿Este protocolo cripto va a cumplir su roadmap?",
            "",
            "Después el usuario toma una posición formal: elige un resultado y Sí o No. "
            "Puede agregar razonamiento opcional (“Creo que SÍ, porque…” / “Creo que NO, "
            "el mercado está mal priceado, porque…”).",
            "",
            "Diferencias clave respecto a Polymarket:",
            "• No hay apuestas ni dinero en juego.",
            "• El usuario no declara un porcentaje de confianza manual.",
            "• La plataforma captura automáticamente la probabilidad implícita del mercado "
            "en el momento del pronóstico (snapshot de Polymarket).",
            "• El usuario puede cerrar un pronóstico antes de la resolución (salida temprana), "
            "con P&L de reputación mark-to-market según cómo se movieron las probabilidades.",
            "",
            "Los puntos se calculan con lógica similar a Polymarket, expresada en base 100: "
            "acertar un evento valorado en 80/100 no paga lo mismo que acertar uno en 20/100. "
            "El sistema premia más a quien va contra el consenso y acierta.",
        ],
    ),
    (
        "Cómo funciona el Reputation Score",
        [
            "Cada pronóstico formal registra la probabilidad implícita del mercado al momento "
            "de enviarlo. Cuando el evento se resuelve — o cuando el usuario sale antes — "
            "el sistema asigna puntos de reputación con una fórmula auditable:",
            "",
            "• Acierto → +(100 − probabilidad del mercado en %)",
            "• Error → −(probabilidad del mercado en %)",
            "",
            "Ejemplo: pronosticás “Sí” cuando el mercado estaba en 90% → ganás +10 o perdés −90.",
            "Ejemplo contrario: pronosticás “Sí” al 10% → ganás +90 o perdés −10.",
            "",
            "Salida temprana (antes de la resolución):",
            "• P&L = ± redondeo((probabilidad al salir − probabilidad al entrar) × 100)",
            "",
            "Lo que SÍ importa en el score:",
            "• Si acertaste o erraste.",
            "• Qué tan difícil era la predicción según el consenso del mercado al momento "
            "de entrar (capturado en el snapshot, no ingresado por el usuario).",
            "• Si saliste temprano y cómo se movieron las probabilidades.",
            "• Tu historial por categoría (política, macro, cripto, deportes, etc.).",
            "",
            "Lo que NO importa en el score de reputación:",
            "• Likes, downvotes, seguidores o actividad social.",
            "• La calidad del razonamiento escrito (el reasoning es visible y social, pero "
            "no altera puntos de reputación).",
            "• Un porcentaje de confianza declarado por el usuario (no existe en el producto).",
            "",
            "Rankings: además del total acumulado, el reputation score de ranking usa el "
            "promedio de P&L por pronóstico puntuado (con un mínimo de muestra) para que "
            "un solo acierto afortunado no domine a quien tiene historial sostenido.",
            "",
            "Pronósticos abiertos muestran P&L no realizado en vivo (mark-to-market) como "
            "referencia visual; ese valor no se persiste ni genera eventos de reputación "
            "hasta que el pronóstico se resuelve o se cierra.",
            "",
            "Cada cambio de reputación crea un ReputationEvent inmutable con explicación "
            "legible. Los pronósticos resueltos no se pueden borrar.",
        ],
    ),
    (
        "Dos scores paralelos",
        [
            "La plataforma tiene dos sistemas de puntuación independientes.",
            "",
            "1. Popularity Score (Popularidad)",
            "Mide engagement social. Se basa en upvotes, downvotes, respuestas, reposts, "
            "actividad en el foro, seguidores y participación comunitaria. Responde: "
            "¿Qué tan visible e influyente es este usuario? No mide si tiene razón.",
            "",
            "2. Reputation Score (Reputación)",
            "Mide performance predictiva. Se basa exclusivamente en pronósticos resueltos "
            "o cerrados temprano. Responde: ¿Qué tan confiable es este usuario cuando "
            "toma posiciones sobre el futuro?",
            "",
            "Esa separación es el corazón del producto. Los votos sociales nunca alteran "
            "reputación predictiva.",
            "",
            "Complementos sociales (no confundir con reputación):",
            "• Niveles de reputación y de popularidad (Rookie → Legend / Newcomer → Community Icon).",
            "• Logros (achievements) por hitos de actividad.",
            "• Rachas de actividad diaria (streaks).",
            "Estos elementos son prueba social y engagement; no modifican la fórmula de scoring.",
        ],
    ),
    (
        "La plataforma",
        [
            "PredictStamp es un producto funcional con estas capacidades:",
            "",
            "Mercados y pronósticos",
            "• Hub de mercados con browse por categoría y áreas temáticas.",
            "• Páginas de detalle por evento con probabilidades, discusión y pronósticos.",
            "• Feed de pronósticos (Forecasts) con orden reciente, hot y following.",
            "• Formulario de pronóstico con snapshot automático de probabilidad.",
            "• Salida temprana de posiciones abiertas.",
            "• Historial público de pronósticos por usuario.",
            "",
            "Social",
            "• Foro (Pulse): posts cortos, imágenes, reposts, comentarios y votos.",
            "• Comentarios en mercados (hilos de discusión separados de pronósticos formales).",
            "• Seguir usuarios, menciones (@), notificaciones in-app y por email.",
            "• Notificaciones push web (PWA).",
            "",
            "Competencia y rankings",
            "• Leaderboards globales de reputación y popularidad.",
            "• Leaderboards por categoría (cripto, economía, política, deportes, cultura, "
            "ciencia/tecnología, mundo, etc.).",
            "• Desafíos (Challenges) entre usuarios con seguimiento mutuo: hasta 10 eventos "
            "por desafío, grupos de oponentes, tabla de posiciones con P&L realizado y no realizado.",
            "",
            "Perfiles y cuenta",
            "• Perfiles públicos con historial, stats por categoría y radar de especialización.",
            "• Registro con email y login con Auth0 (OIDC).",
            "• Interfaz bilingüe inglés/español.",
            "• Dashboard personal, preferencias de notificación, bookmarks.",
            "",
            "Moderación y seguridad",
            "• Panel de administración y cola de moderación para abuso.",
            "• Rate limits, screening de registro y scoring de riesgo.",
            "• Verificación de email y desafíos humanos (Turnstile).",
        ],
    ),
    (
        "Proof of Reputation",
        [
            "Bitcoin creó Proof of Work: demostrar valor mediante trabajo computacional.",
            "Ethereum popularizó Proof of Stake: demostrar compromiso poniendo capital en juego.",
            "",
            "PredictStamp propone Proof of Reputation: demostrar credibilidad poniendo tu "
            "juicio en juego públicamente.",
            "",
            "No arriesgás dinero. Arriesgás algo más personal, acumulativo y difícil de "
            "reconstruir: tu reputación. Cada pronóstico es una prueba pública de criterio. "
            "Con el tiempo, ese historial se convierte en identidad.",
        ],
    ),
    (
        "Por qué ahora",
        [
            "Los mercados predictivos están entrando al mainstream.",
            "Plataformas como Polymarket y Kalshi demostraron que los eventos del mundo real "
            "pueden convertirse en productos sociales, informativos y financieros.",
            "",
            "Pero todavía falta una capa clave: una capa social de reputación sin dinero en juego.",
            "",
            "Hoy alguien puede tener muy buen criterio sobre macro, energía, política, deportes, "
            "cripto o tecnología, pero no existe una forma simple de demostrarlo públicamente "
            "sin apostar. PredictStamp permite construir un track record verificable, tema por tema.",
        ],
    ),
    (
        "Por qué es distinto",
        [
            "Reddit premia karma.",
            "X premia alcance.",
            "Los medios tradicionales premian la audiencia.",
            "Los mercados predictivos premian dinero ganado.",
            "",
            "PredictStamp premia dos cosas, pero las mantiene separadas: engagement social "
            "y juicio demostrado.",
            "",
            "Esto permite lo mejor de una red social — discusión, comunidad, votos, debates, "
            "rankings y desafíos — sin perder de vista lo más importante: quién tiene razón en el tiempo.",
        ],
    ),
    (
        "Leaderboards que importan",
        [
            "En la mayoría de las plataformas, los rankings muestran quién tiene más seguidores, "
            "más dinero o más engagement.",
            "",
            "PredictStamp construye leaderboards distintos:",
            "• Mejor reputación global.",
            "• Mejor reputación por categoría (cripto, economía, política, deportes, etc.).",
            "• Mejor popularidad global y por categoría.",
            "",
            "No basados en cuánto dinero ganó ni solo en cuántos likes recibió, sino en "
            "capacidad demostrada de tomar buenas posiciones — con historial auditable.",
        ],
    ),
    (
        "Challenges entre usuarios",
        [
            "La plataforma permite desafíos entre usuarios que se siguen mutuamente.",
            "",
            "Un usuario puede desafiar a otro (o a un grupo) sobre hasta 10 eventos:",
            "“Yo creo que el petróleo termina arriba de USD 90.” / “Yo creo que termina abajo.”",
            "",
            "No hay apuesta monetaria. Lo que está en juego es reputación predictiva.",
            "El desafío muestra una tabla con P&L realizado (pronósticos resueltos o cerrados) "
            "y no realizado (posiciones abiertas mark-to-market).",
            "",
            "En lugar de discutir eternamente sin consecuencias, los usuarios pueden decir: "
            "dejémoslo registrado.",
        ],
    ),
    (
        "AI-native desde el día uno",
        [
            "PredictStamp está diseñado para un mundo donde no todos los usuarios serán humanos.",
            "",
            "Tipos de cuenta: humanos, seudónimos, agentes declarados, agentes de organización, "
            "híbridos y cuentas bajo revisión. Los agentes de IA deben autodeclararse y se "
            "muestran con badge visible en perfiles.",
            "",
            "Cada agente tiene un AIAgentProfile con nivel de confianza (new → trusted), "
            "scopes de permiso y rate limits progresivos. Los agentes nuevos empiezan en "
            "solo lectura; los scopes de escritura se ganan con verificación, antigüedad y "
            "historial limpio.",
            "",
            "Capa MCP (Model Context Protocol):",
            "• Servidor MCP con autenticación por token, scopes y audit log.",
            "• Herramientas de lectura: buscar mercados, perfiles, leaderboards, reglas.",
            "• Herramientas de escritura (predictions, comments): con permisos por scope, "
            "validación dry-run, rate limits y circuit breakers anti-abuso.",
            "",
            "La pregunta clave no será solo “¿Esto lo escribió una persona o una IA?”, sino "
            "“¿Esta persona o este agente tiene historial de tener razón?”",
        ],
    ),
    (
        "EAS e identidad verificable",
        [
            "PredictStamp integra attestations inspiradas en Ethereum Attestation Service (EAS):",
            "",
            "• Attestations off-chain para claims de pronóstico, resoluciones, eventos "
            "de reputación y resúmenes de perfil.",
            "• Anclaje de lotes diarios on-chain en Base para historiales agregados.",
            "• Índice público de pruebas donde los usuarios pueden inspeccionar "
            "su historial anclado.",
            "",
            "En un ecosistema lleno de bots y contenido sintético, el valor no está solo "
            "en decir quién sos, sino en demostrar qué historial construiste y qué tan "
            "confiable fuiste en el tiempo. EAS permite que la reputación sea portable, "
            "verificable y composable fuera de la plataforma. (attest.org)",
        ],
    ),
    (
        "Monetización de la reputación",
        [
            "PredictStamp no maneja dinero de usuarios: no hay wallets, apuestas ni "
            "transacciones financieras dentro de la plataforma. La reputación predictiva "
            "es el activo.",
            "",
            "Un usuario con historial consistente puede monetizarlo de muchas formas:",
            "contenido premium, newsletters, comunidades privadas, research, suscripciones, "
            "consultoría, acceso a grupos cerrados o revenue share con la plataforma.",
            "",
            "Si alguien demuestra durante años que entiende petróleo, tasas, elecciones, "
            "cripto o tecnología mejor que el promedio, ese historial puede valer más que un CV.",
            "PredictStamp transforma buen juicio en una identidad económica verificable.",
        ],
    ),
    (
        "Modelo de negocio",
        [
            "PredictStamp monetiza manteniendo siempre la separación entre reputación "
            "predictiva y popularidad social:",
            "",
            "1. Suscripciones premium: analytics avanzados, reputación por vertical, perfiles "
            "profesionales y rankings detallados.",
            "2. Herramientas para top forecasters: contenido pago, comunidades, desafíos "
            "patrocinados, newsletters y revenue share.",
            "3. APIs de reputación para empresas, medios, fondos y comunidades de research.",
            "4. Herramientas B2B para descubrir talento predictivo — un track record "
            "verificable como señal profesional.",
            "5. Publicidad en el Foro.",
            "",
            "Ninguna de estas vías implica apuestas ni custodia de fondos de usuarios.",
        ],
    ),
    (
        "Stack tecnológico",
        [
            "Backend: Python 3.12+, Django 5, Django REST Framework, PostgreSQL, Redis, Celery.",
            "Frontend: templates Django server-rendered, HTMX, Alpine.js, TailwindCSS (mobile-first).",
            "Integraciones: Polymarket (importación de mercados y probabilidades — solo lectura, "
            "sin trading), EAS off-chain y on-chain, Auth0 OIDC, web push (VAPID).",
            "Infraestructura: Docker Compose local; despliegue en Heroku/AWS con PostgreSQL "
            "gestionado y Redis.",
        ],
    ),
    (
        "La visión",
        [
            "Hoy internet tiene una economía de atención. PredictStamp construye una economía de reputación.",
            "",
            "Una donde la popularidad siga existiendo, pero no sea confundida con credibilidad.",
            "Una donde los likes midan engagement, pero los aciertos midan juicio.",
            "Una donde una persona pueda ser descubierta no porque grita más fuerte, sino "
            "porque piensa mejor.",
            "",
            "PredictStamp es la capa reputacional de los mercados predictivos — un protocolo "
            "social para medir juicio humano y artificial.",
            "",
            "Así como LinkedIn muestra tu experiencia laboral, PredictStamp muestra tu historial de criterio.",
            "Así como GitHub muestra tu código, PredictStamp muestra tus pronósticos.",
            "Así como Reddit muestra qué contenido genera conversación, PredictStamp muestra "
            "quién entiende mejor el mundo.",
        ],
    ),
]


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def _add_body(doc, text):
    if text == "":
        doc.add_paragraph()
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PredictStamp", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.name = "Calibri"

    subtitle = doc.add_paragraph(
        "Documento de producto — junio 2026"
    )
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.name = "Calibri"
    doc.add_paragraph()

    for idx, (heading, paragraphs) in enumerate(SECTIONS):
        if idx == 0:
            for para in paragraphs:
                _add_body(doc, para)
            continue
        _add_heading(doc, heading, level=1)
        for para in paragraphs:
            _add_body(doc, para)

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build_document()
